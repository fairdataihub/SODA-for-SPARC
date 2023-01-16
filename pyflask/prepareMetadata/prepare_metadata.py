# -*- coding: utf-8 -*-

### Import required python modules

from gevent import monkey

monkey.patch_all()
import platform
import os
from os.path import (
    isdir,
    join,
    exists,
    expanduser,
    getsize,
)
import pandas as pd
import csv
import shutil
import numpy as np
from collections import defaultdict
from pennsieve import Pennsieve
import requests
from errorHandlers import is_file_not_found_exception, is_invalid_file_exception, InvalidDeliverablesDocument

from string import ascii_uppercase
import itertools
import tempfile

from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font
from docx import Document

from flask import abort 

from manageDatasets import (
    bf_get_current_user_permission,
)
from curate import create_high_level_manifest_files_existing_bf_starting_point, get_name_extension

from pysodaUtils import agent_running

from namespaces import NamespaceEnum, get_namespace_logger
namespace_logger = get_namespace_logger(NamespaceEnum.CURATE_DATASETS)

userpath = expanduser("~")
METADATA_UPLOAD_BF_PATH = join(userpath, "SODA", "METADATA")
TEMPLATE_PATH = ""

### Sets the TEMPLATE_PATH using SODA-for-SPARC's basepath so that the prepare_metadata section can find
### the templates stored in file_templates direcotory
def set_template_path(soda_base_path, soda_resources_path):
    """
    Sets the TEMPLATE_PATH using SODA-for-SPARC's basepath so that the prepare_metadata section can find
    the templates stored in file_templates direcotory.
    """
    global TEMPLATE_PATH


    # once pysoda has been packaged with pyinstaller
    # it creates an archive that slef extracts to an OS-specific temp directory.
    # Due to this we can no longer use a relative path from the pysoda directory to the file_templates folder.
    # When running in dev mode this also works
    TEMPLATE_PATH = join(soda_base_path, "file_templates")

    # check if os is Darwin/Linux
    if platform.system() in ["Darwin", "Linux"] and not exists(TEMPLATE_PATH):
        # we are in production and we need to use the Resources folder for the file_templates folder
        TEMPLATE_PATH = join(soda_resources_path, "file_templates")





### Import Data Deliverables document
def import_milestone(filepath):
    """
        Import Data Deliverables document
    """
    doc = Document(filepath)
    try:
        table = doc.tables[0]
    except IndexError:
        raise InvalidDeliverablesDocument(
            "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: 'Related milestone, aim, or task', 'Description of data', and 'Expected date of completion'."
        )
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # headers will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def extract_milestone_info(datalist):
    milestone = defaultdict(list)
    milestone_key1 = "Related milestone, aim, or task"
    milestone_key2 = "Related milestone, aim or task"
    other_keys = ["Description of data", "Expected date of completion"]
    for row in datalist:
        if milestone_key1 in row:
            milestone_key = milestone_key1
        elif milestone_key2 in row:
            milestone_key = milestone_key2
        else:
            raise InvalidDeliverablesDocument(
                "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: Related milestone, aim, or task, Description of data, and Expected date of completion."
            )

        key = row[milestone_key]
        if key != "":
            milestone[key].append({key: row[key] for key in other_keys})

    return milestone


### Create submission file
def save_submission_file(upload_boolean, bfaccount, bfdataset, filepath, val_arr):

    font_submission = Font(name="Calibri", size=14, bold=False)

    source = join(TEMPLATE_PATH, "submission.xlsx")

    destination = join(METADATA_UPLOAD_BF_PATH, "submission.xlsx") if upload_boolean else filepath

    try:
        shutil.copyfile(source, destination)
    except FileNotFoundError as e:
        raise e

    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]
    for column, arr in zip(excel_columns(start_index=2), val_arr):
        ws1[column + "2"] = arr["award"]
        ws1[column + "3"] = arr["milestone"]
        ws1[column + "4"] = arr["date"]

        ws1[column + "2"].font = font_submission
        ws1[column + "3"].font = font_submission
        ws1[column + "4"].font = font_submission

    rename_headers(ws1, len(val_arr), 2)

    wb.save(destination)

    # calculate the size of the metadata file
    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file("submission.xlsx", bfaccount, bfdataset, destination, True)

    return {"size": size}


# this function saves and uploads the README/CHANGES to Pennsieve, just when users choose to generate onto Pennsieve
## (not used for generating locally)
def upload_RC_file(text_string, file_type, bfaccount, bfdataset):

    file_path = join(METADATA_UPLOAD_BF_PATH, file_type)

    with open(file_path, "w") as f:
        f.write(text_string)

    size = getsize(file_path)

    upload_metadata_file(file_type, bfaccount, bfdataset, file_path, True)

    return { "size": size, "filepath": file_path }


def upload_metadata_file(file_type, bfaccount, bfdataset, file_path, delete_after_upload):
    ## check if agent is running in the background
    agent_running()

    try:
        bf = Pennsieve(bfaccount)
    except Exception:
        abort(400, "Please select a valid Pennsieve account.")
    
    # check that the Pennsieve dataset is valid
    try:
        myds = bf.get_dataset(bfdataset)
    except Exception:
        abort(400, "Please select a valid Pennsieve dataset.")


    # check that the user has permissions for uploading and modifying the dataset
    role = bf_get_current_user_permission(bf, myds)
    if role not in ["owner", "manager", "editor"]:
        abort(403, "You don't have permissions for uploading to this Pennsieve dataset.")

    # handle duplicates on Pennsieve: first, obtain the existing file ID
    for i in range(len(myds.items)):

        if myds.items[i].name == file_type:

            item_id = myds.items[i].id

            # then, delete it using Pennsieve method delete(id)
            bf.delete(item_id)

    myds.upload(file_path)
    # delete the local file that was created for the purpose of uploading to Pennsieve
    if delete_after_upload:
        os.remove(file_path)

def excel_columns(start_index=0):
    """
    NOTE: does not support more than 699 contributors/links
    """
    single_letter = list(ascii_uppercase[start_index:])
    two_letter = [a + b for a, b in itertools.product(ascii_uppercase, ascii_uppercase)]
    return single_letter + two_letter
  

def rename_headers(workbook, max_len, start_index):
    """
    Rename header columns if values exceed 3. Change Additional Values to Value 4, 5,...
    """

    columns_list = excel_columns(start_index=start_index)
    if max_len >= start_index:

        workbook[columns_list[0] + "1"] = "Value"

        for i, column in zip(range(2, max_len + 1), columns_list[1:]):

            workbook[column + "1"] = f"Value {str(i)}"
            cell = workbook[column + "1"]

            blueFill = PatternFill(
                start_color="9CC2E5", end_color="9CC2E5", fill_type="solid"
            )

            font = Font(bold=True)
            cell.fill = blueFill
            cell.font = font

    else:

        delete_range = len(columns_list) - max_len
        workbook.delete_cols(4 + max_len, delete_range)


def grayout_subheaders(workbook, max_len, start_index):
    """
    Gray out sub-header rows for values exceeding 3 (SDS2.0).
    """
    headers_list = ["4", "10", "18", "23", "28"]
    columns_list = excel_columns(start_index=start_index)

    for (i, column), no in itertools.product(zip(range(2, max_len + 1), columns_list[1:]), headers_list):
        cell = workbook[column + no]
        fillColor("B2B2B2", cell)


def grayout_single_value_rows(workbook, max_len, start_index):
    """
    Gray out rows where only single values are allowed. Row number: 2, 3, 5, 6, 9, 11, 12, 13, 17, 29, 30
    """

    columns_list = excel_columns(start_index=start_index)
    row_list = ["2", "3", "5", "6", "9", "11", "12", "13", "17", "29", "30"]
    for (i, column), no in itertools.product(zip(range(2, max_len + 1), columns_list[1:]), row_list):
        cell = workbook[column + no]
        fillColor("CCCCCC", cell)


def fillColor(color, cell):
    colorFill = PatternFill(start_color=color, end_color=color, fill_type="solid")

    cell.fill = colorFill


### Prepare dataset-description file


def populate_dataset_info(ws, val_obj):
    ## name, description, type, samples, subjects
    ws["D5"] = val_obj["name"]
    ws["D6"] = val_obj["description"]
    ws["D3"] = val_obj["type"]
    ws["D29"] = val_obj["number of subjects"]
    ws["D30"] = val_obj["number of samples"]

    ## keywords
    for i, column in zip(range(len(val_obj["keywords"])), excel_columns(start_index=3)):
        ws[column + "7"] = val_obj["keywords"][i]

    return val_obj["keywords"]


def populate_study_info(workbook, val_obj):
    workbook["D11"] = val_obj["study purpose"]
    workbook["D12"] = val_obj["study data collection"]
    workbook["D13"] = val_obj["study primary conclusion"]
    workbook["D17"] = val_obj["study collection title"]

    ## study organ system
    for i, column in zip(
        range(len(val_obj["study organ system"])), excel_columns(start_index=3)
    ):
        workbook[column + "14"] = val_obj["study organ system"][i]
    ## study approach
    for i, column in zip(
        range(len(val_obj["study approach"])), excel_columns(start_index=3)
    ):
        workbook[column + "15"] = val_obj["study approach"][i]
    ## study technique
    for i, column in zip(
        range(len(val_obj["study technique"])), excel_columns(start_index=3)
    ):
        workbook[column + "16"] = val_obj["study technique"][i]

    return max(
        len(val_obj["study organ system"]),
        len(val_obj["study approach"]),
        len(val_obj["study technique"]),
    )


def populate_contributor_info(workbook, val_array):
    ## award info
    for i, column in zip(
        range(len(val_array["funding"])), excel_columns(start_index=3)
    ):
        workbook[column + "8"] = val_array["funding"][i]

    ### Acknowledgments
    workbook["D9"] = val_array["acknowledgment"]

    ### Contributors
    for contributor, column in zip(
        val_array["contributors"], excel_columns(start_index=3)
    ):
        workbook[column + "19"] = contributor["conName"]
        workbook[column + "20"] = contributor["conID"]
        workbook[column + "21"] = contributor["conAffliation"]
        workbook[column + "22"] = contributor["conRole"]

    return [val_array["funding"], val_array["contributors"]]


def populate_related_info(workbook, val_array):
    ## related links including protocols

    for i, column in zip(range(len(val_array)), excel_columns(start_index=3)):
        workbook[column + "24"] = val_array[i]["description"]
        workbook[column + "25"] = val_array[i]["relation"]
        workbook[column + "26"] = val_array[i]["link"]
        workbook[column + "27"] = val_array[i]["type"]

    return len(val_array)


### generate the dataset_description file


def save_ds_description_file(
    upload_boolean,
    bfaccount,
    bfdataset,
    filepath,
    dataset_str,
    study_str,
    con_str,
    related_info_str,
):
    source = join(TEMPLATE_PATH, "dataset_description.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "dataset_description.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)

    # json array to python list
    val_obj_study = study_str
    val_obj_ds = dataset_str
    val_arr_con = con_str
    val_arr_related_info = related_info_str

    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    ws1["D22"] = ""
    ws1["E22"] = ""
    ws1["D24"] = ""
    ws1["E24"] = ""
    ws1["D25"] = ""
    ws1["E25"] = ""

    keyword_array = populate_dataset_info(ws1, val_obj_ds)

    study_array_len = populate_study_info(ws1, val_obj_study)

    (funding_array, contributor_role_array) = populate_contributor_info(
        ws1, val_arr_con
    )

    related_info_len = populate_related_info(ws1, val_arr_related_info)

    # keywords length
    keyword_len = len(keyword_array)

    # contributors length
    no_contributors = len(contributor_role_array)

    # funding = SPARC award + other funding sources
    funding_len = len(funding_array)

    # obtain length for formatting compliance purpose
    max_len = max(
        keyword_len, funding_len, no_contributors, related_info_len, study_array_len
    )

    rename_headers(ws1, max_len, 3)
    grayout_subheaders(ws1, max_len, 3)
    grayout_single_value_rows(ws1, max_len, 3)

    if ws1["G1"].value == "Value n":
        ws1.delete_cols(7)

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file(
            "dataset_description.xlsx", bfaccount, bfdataset, destination, True
        )

    return {"size": size}


subjectsTemplateHeaderList = [
    "subject id",
    "pool id",
    "subject experimental group",
    "age",
    "sex",
    "species",
    "strain",
    "rrid for strain",
    "age category",
    "also in dataset",
    "member of",
    "laboratory internal id",
    "date of birth",
    "age range (min)",
    "age range (max)",
    "body mass",
    "genotype",
    "phenotype",
    "handedness",
    "reference atlas",
    "experimental log file path",
    "experiment date",
    "disease or disorder",
    "intervention",
    "disease model",
    "protocol title",
    "protocol url or doi",
]
samplesTemplateHeaderList = [
    "sample id",
    "subject id",
    "was derived from",
    "pool id",
    "sample experimental group",
    "sample type",
    "sample anatomical location",
    "also in dataset",
    "member of",
    "laboratory internal id",
    "date of derivation",
    "experimental log file path",
    "reference atlas",
    "pathology",
    "laterality",
    "cell type",
    "plane of section",
    "protocol title",
    "protocol url or doi",
]

def upload_code_description_metadata(filepath, bfAccount, bfDataset):
    upload_metadata_file("code_description.xlsx", bfAccount, bfDataset, filepath, False)
        


def save_subjects_file(upload_boolean, bfaccount, bfdataset, filepath, datastructure):

    source = join(TEMPLATE_PATH, "subjects.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "subjects.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:11]
    optionalFields = transposeDatastructure[11:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = subjectsTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(
            np.concatenate((sortMatrix, refinedOptionalFields))
        )
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)
    #
    # # 1. delete rows using delete_rows(index, amount=2) -- description and example rows
    # ws1.delete_rows(2, 2)
    # delete all optional columns first (from the template)
    ws1.delete_cols(12, 18)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(
        start_color="FFD965", end_color="FFD965", fill_type="solid"
    )

    # gevent.sleep(0)
    for column, header in zip(
        excel_columns(start_index=11), refinedDatastructure[0][11:headers_no]
    ):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name="Calibri")

    # gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            ws1[cell] = refinedDatastructure[i][j] or ""
            ws1[cell].font = Font(bold=False, size=11, name="Arial")

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file("subjects.xlsx", bfaccount, bfdataset, destination, True)

    return size


def save_samples_file(upload_boolean, bfaccount, bfdataset, filepath, datastructure):
    source = join(TEMPLATE_PATH, "samples.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "samples.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)

    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:9]
    optionalFields = transposeDatastructure[9:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = samplesTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(
            np.concatenate((sortMatrix, refinedOptionalFields))
        )
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)

    ws1.delete_cols(10, 15)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(
        start_color="FFD965", end_color="FFD965", fill_type="solid"
    )
    # gevent.sleep(0)
    for column, header in zip(
        excel_columns(start_index=9), refinedDatastructure[0][9:headers_no]
    ):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name="Calibri")

    # gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            ws1[cell] = refinedDatastructure[i][j] or ""
            ws1[cell].font = Font(bold=False, size=11, name="Arial")

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, call upload function
    if upload_boolean:
        upload_metadata_file("samples.xlsx", bfaccount, bfdataset, destination, True)

    return {"size": size}


# check for non-empty fields (cells)
def column_check(x):
    return "unnamed" not in x.lower()


# import an existing subjects/samples files from an excel file
def convert_subjects_samples_file_to_df(type, filepath, ui_fields):

    subjects_df = pd.read_excel(
        filepath, engine="openpyxl", usecols=column_check, header=0
    )
    subjects_df = subjects_df.dropna(axis=0, how="all")
    subjects_df = subjects_df.replace(np.nan, "", regex=True)
    subjects_df = subjects_df.applymap(str)
    subjects_df.columns = map(str.lower, subjects_df.columns)

    if type == "subjects":
        if "subject id" not in list(subjects_df.columns.values):
            abort(
                400, 
                "The header 'subject id' is required to import an existing subjects file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/subjects.xlsx'>template</a> of the subjects file."
            )

        elif checkEmptyColumn(subjects_df["subject id"]):
            abort(
                400,
                "At least 1 'subject id' is required to import an existing subjects file"
            )

        templateHeaderList = subjectsTemplateHeaderList

    else:
        if "subject id" not in list(subjects_df.columns.values) or "sample id" not in list(subjects_df.columns.values):
            abort(
                400,
                "The headers 'subject id' and 'sample id' are required to import an existing samples file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/samples.xlsx'>template</a> of the samples file."
            )

        if checkEmptyColumn(subjects_df["sample id"]) or checkEmptyColumn(
            subjects_df["sample id"]
        ):
            abort(
                400,
                "At least 1 'subject id' and 'sample id' pair is required to import an existing samples file"
            )

        templateHeaderList = samplesTemplateHeaderList

    importedHeaderList = list(subjects_df.columns.values)
    print(importedHeaderList)
    print("#" * 30)

    transpose = []
    for header in templateHeaderList:
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    # print(header)
    for header in importedHeaderList:
        print("$" * 40)
        print(header)
        print("$" * 40)

        if header.lower() in templateHeaderList:
            continue
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    print(transpose)
    print("#" * 30)
    print(ui_fields)
    sortMatrix = sortedSubjectsTableData(transpose, ui_fields)

    return {"sample_file_rows": transposeMatrix(sortMatrix)} if type in ["samples.xlsx", "samples"] else {"subject_file_rows": transposeMatrix(sortMatrix)}


### function to read existing Pennsieve manifest files and load info into a dictionary
def convert_manifest_to_dict(url):

    manifest_df = pd.read_excel(url, engine="openpyxl", usecols=column_check, header=0)
    manifest_df = manifest_df.dropna(axis=0, how="all")
    manifest_df = manifest_df.replace(np.nan, "", regex=True)
    manifest_df = manifest_df.applymap(str)

    return manifest_df.to_dict()


def checkEmptyColumn(column):
    for element in column:
        if element:
            break
        return True
    return False


# needed to sort subjects and samples table data to match the UI fields
def sortedSubjectsTableData(matrix, fields):
    sortedMatrix = []
    for field in fields:
        for column in matrix:
            if column[0].lower() == field:
                sortedMatrix.append(column)
                break

    customHeaderMatrix = [
        column for column in matrix if column[0].lower() not in fields
    ]

    return (
        np.concatenate((sortedMatrix, customHeaderMatrix)).tolist()
        if customHeaderMatrix
        else sortedMatrix
    )


# transpose a matrix (array of arrays)
def transposeMatrix(matrix):
    return [[matrix[j][i] for j in range(len(matrix))] for i in range(len(matrix[0]))]


# helper function to process custom fields (users add and name them) for subjects and samples files
def processMetadataCustomFields(matrix):
    return [column for column in matrix if any(column[1:])]


## check if any whole column from Excel sheet is empty
def checkEmptyColumn(column):
    for element in column:
        if element:
            break
        return True
    return False



## load/import an existing local or Pennsieve submission.xlsx file
def load_existing_submission_file(filepath):

    try:
        DD_df = pd.read_excel(
            filepath, engine="openpyxl", usecols=column_check, header=0
        )

    except Exception as e:
        if is_file_not_found_exception(e):
            abort(400, "Local submission file not found")

        # TODO: TEST check if error can indicate if file is not in the correct format TEST
        if is_invalid_file_exception(e):
            abort(400, "Local submission file is not in the correct format")

        raise Exception(
            "SODA cannot read this submission.xlsx file. If you are trying to retrieve a submission.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
        ) from e

    DD_df = DD_df.dropna(axis=0, how="all")
    DD_df = DD_df.replace(np.nan, "", regex=True)
    DD_df = DD_df.applymap(str)
    DD_df = DD_df.applymap(str.strip)

    basicColumns = ["Submission Item", "Definition", "Value"]
    basicHeaders = [
        "SPARC Award number",
        "Milestone achieved",
        "Milestone completion date",
    ]
    ## normalize the entries to lowercase just for Version Exception check
    basicColumns = [x.lower() for x in basicColumns]
    basicHeaders = [x.lower() for x in basicHeaders]
    DD_df_lower = [x.lower() for x in DD_df]

    for key in basicColumns:
        if key not in DD_df_lower:
            abort(
                400,
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )

    for header_name in basicHeaders:
        submissionItems = [x.lower() for x in DD_df["Submission Item"]]
        if header_name not in set(submissionItems):
            abort(
                400,
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )

    awardNo = DD_df["Value"][0]
    milestones = [DD_df["Value"][1]]

    for i in range(3, len(DD_df.columns)):
        value = DD_df[f"Value {str(i - 1)}"]
        milestones.append(value[1])

    date = DD_df["Value"][2] or ""

    return {
        "SPARC Award number": awardNo,
        "Milestone achieved": milestones,
        "Milestone completion date": date,
    }


# import existing metadata files except Readme and Changes from Pennsieve
def import_bf_metadata_file(file_type, ui_fields, bfaccount, bfdataset):
    try: 
        bf = Pennsieve(bfaccount)
    except Exception:
        abort(400, "Please select a valid Pennsieve account.")

    try: 
        myds = bf.get_dataset(bfdataset)
    except Exception:
        abort(400, "Please select a valid Pennsieve dataset.")
    

    for i in range(len(myds.items)):

        if myds.items[i].name == file_type:

            item_id = myds.items[i].id
            url = returnFileURL(bf, item_id)

            if file_type == "submission.xlsx":
                return load_existing_submission_file(url)

            elif file_type == "dataset_description.xlsx":
                return load_existing_DD_file("bf", url)

            elif file_type == "subjects.xlsx":
                return convert_subjects_samples_file_to_df("subjects", url, ui_fields)

            elif file_type == "samples.xlsx":
                return convert_subjects_samples_file_to_df("samples", url, ui_fields)

    abort(400, 
        f"No {file_type} file was found at the root of the dataset provided."
    )


# import readme or changes file from Pennsieve
def import_bf_RC(bfaccount, bfdataset, file_type):

    file_type = file_type + ".txt"

    try:
        bf = Pennsieve(bfaccount)
    except Exception:
        abort(400, "Please select a valid Pennsieve account.")

    try:
        myds = bf.get_dataset(bfdataset)
    except Exception:
        abort(400, "Please select a valid Pennsieve dataset.")

    for i in range(len(myds.items)):

        if myds.items[i].name == file_type:

            item_id = myds.items[i].id
            url = returnFileURL(bf, item_id)

            response = requests.get(url)
            return {"text": response.text}

    abort (400, f"No {file_type} file was found at the root of the dataset provided.")


# path to local SODA folder for saving manifest files
manifest_folder_path = join(userpath, "SODA", "manifest_files")
manifest_progress = {
    "total_manifest_files": 0,
    "manifest_files_uploaded": 0,
    "finished": False
}

# TODO: NOTE: ESSENTIAL: Remove the manifest_file even if the user does not generate before pulling again.
def import_bf_manifest_file(soda_json_structure, bfaccount, bfdataset):
    # reset the progress tracking information
    global manifest_progress
    manifest_progress["finished"] = False
    manifest_progress["total_manifest_files"] = 0
    manifest_progress["manifest_files_uploaded"] = 0

    bf = Pennsieve(bfaccount)
    myds = bf.get_dataset(bfdataset)

    high_level_folders = ["code", "derivative", "docs", "primary", "protocol", "source"]

    dataset_structure = soda_json_structure["dataset-structure"]

    # get the count of the total number of high level folders in soda_json_structure
    for folder in list(dataset_structure["folders"]):
        if folder in high_level_folders:
            manifest_progress["total_manifest_files"] += 1

    # create the path to the dataset files and folders on Pennsieve and add them to the dataset structure stored in soda_json_structure
    recursive_item_path_create(dataset_structure, [])

    high_level_folders = ["code", "derivative", "docs", "primary", "protocol", "source"]

    # handle updating any existing manifest files on Pennsieve
    update_existing_pennsieve_manifest_files(myds, bf, dataset_structure, high_level_folders)

    # create manifest files from scratch for any high level folders that don't have a manifest file on Pennsieve
    create_high_level_manifest_files_existing_bf_starting_point(soda_json_structure, high_level_folders, manifest_progress)

    # finished with the manifest generation process
    manifest_progress["finished"] = True

    no_manifest_boolean = False


def update_existing_pennsieve_manifest_files(myds, bf, dataset_structure, high_level_folders):
    global manifest_progress
    # handle updating any existing manifest files on Pennsieve
    for i in range(len(myds.items)):
        if myds.items[i].name in [
            "code",
            "derivative",
            "docs",
            "primary",
            "protocol",
            "source",
        ]:
            for j in range(len(myds.items[i])):
                if myds.items[i][j].name == "manifest.xlsx":

                    if not exists(join(manifest_folder_path, myds.items[i].name)):
                        # create the path
                        os.makedirs(join(manifest_folder_path, myds.items[i].name))

                    item_id = myds.items[i][j].id
                    url = returnFileURL(bf, item_id)

                    manifest_df = pd.read_excel(
                        url, engine="openpyxl", usecols=column_check, header=0
                    )

                    filepath = join(
                        manifest_folder_path, myds.items[i].name, "manifest.xlsx"
                    )

                    high_level_folders.remove(myds.items[i].name)

                    updated_manifest_dict = update_existing_pennsieve_manifest_file(dataset_structure["folders"][myds.items[i].name], manifest_df)

                    if not exists(join(manifest_folder_path, myds.items[i].name)):
                        # create the path
                        os.makedirs(join(manifest_folder_path, myds.items[i].name))

                    new_manifest = pd.DataFrame.from_dict(updated_manifest_dict)
                    new_manifest.to_excel(filepath, index=False)
                    wb = load_workbook(filepath)
                    ws = wb.active
                    blueFill = PatternFill(
                        start_color="9DC3E6", fill_type="solid"
                    )
                    greenFill = PatternFill(
                        start_color="A8D08D", fill_type="solid"
                    )
                    yellowFill = PatternFill(
                        start_color="FFD965", fill_type="solid"
                    )
                    ws['A1'].fill = blueFill
                    ws['B1'].fill = greenFill
                    ws['C1'].fill = greenFill
                    ws['D1'].fill = greenFill
                    ws['E1'].fill = yellowFill
                    wb.save(filepath)

                    manifest_progress["manifest_files_uploaded"] += 1

                    no_manifest_boolean = True

                    # break because we only need to read the "manifest.xlsx" file in each high level folder.
                    break


def update_existing_pennsieve_manifest_file(high_level_folder, manifest_df):
    """
        Given a high level folder and the existing manifest file therein, create a new updated manifest file. 
        This manifest file needs to have files that only exist in the high level folder. Additionally, it needs to 
        retain any custom metadata that the user has added to the old manifest file for entries that still exist. 
    """

    new_manifest_dict = {'filename': [], 'timestamp': [], 'description': [], 'file type': [], 'Additional Metadata': []}
    SET_COLUMNS = ['filename', 'timestamp', 'description', 'file type', 'Additional Metadata']
    for column in manifest_df.columns: 
        if column not in new_manifest_dict:
            new_manifest_dict[column] = []
            SET_COLUMNS.append(column)

    # convert the old manifest into a dictionary to optimize the lookup time
    old_manifest_dict = {x:manifest_df[x].values.tolist() for x in manifest_df}

    # create a mapping of filename to the idx of the row in the old_manidest_dict
    filename_idx_map = {x:i for i, x in enumerate(manifest_df['filename'])}

    # traverse through the high level folder items
    update_existing_pennsieve_manifest_file_helper(high_level_folder, old_manifest_dict, new_manifest_dict, filename_idx_map, manifest_columns=SET_COLUMNS)

    # write the new manifest_df to the correct location as an excel file 
    return new_manifest_dict


def update_existing_pennsieve_manifest_file_helper(folder, old_manifest_dict, new_manifest_dict, filename_idx_map, manifest_columns):
    """
        Traverse through each high level folder in the dataset and update the new manifest data frame.
    """

    if "files" in folder.keys():
        for file in list(folder["files"]):
            file_path = remove_high_level_folder_from_path(folder["files"][file]["folderpath"]) + f"{file}"

            # select the row in the old manifest file that has the same file path as the file in the current folder
            # rationale: this means the file still exists in the user's dataset
            row_idx = filename_idx_map.get(file_path, None)

            if row_idx is None:
                for key in new_manifest_dict.keys():
                    if key == "filename":
                        new_manifest_dict["filename"].append(file_path)   
                    elif key == "timestamp":
                        new_manifest_dict["timestamp"].append(folder["files"][file]["timestamp"]), 
                    elif key == "description": 
                        new_manifest_dict["description"].append(folder["files"][file].get("description", ""))
                    elif key == "file type":
                        unused_file_name, file_extension = get_name_extension(file)
                        new_manifest_dict["file type"].append(file_extension), 
                    elif key == "Additional Metadata":
                        new_manifest_dict["Additional Metadata"].append(folder["files"][file].get("additional-metadata", ""))
                    else:
                        new_manifest_dict[key].append("")
                    

            else:
                # add the existing rows to the new manifest dictionary's arrays
                # TODO: Confirm it adds NULL/NaN if the value is empty
                for column in manifest_columns:
                    new_manifest_dict[column].append(old_manifest_dict[column][row_idx])

    if "folders" in folder.keys():
        for current_folder in list(folder["folders"]):
            update_existing_pennsieve_manifest_file_helper(folder["folders"][current_folder], old_manifest_dict, new_manifest_dict, filename_idx_map, manifest_columns)

    return 


def manifest_creation_progress():
    """
        Monitors the progress of the manifest creation process.
    """
    global manifest_progress

    return {
            "manifest_files_uploaded": manifest_progress["manifest_files_uploaded"], 
            "total_manifest_files": manifest_progress["total_manifest_files"],
            "finished": manifest_progress["finished"]
           }


def remove_high_level_folder_from_path(paths):
    """
        Remove the high level folder from the path. This is necessary because the high level folder is not included in the manifest file name entry.
    """

    return "" if len(paths) == 1 else "/".join(paths[1:]) + "/"




def copytree(src, dst, symlinks=False, ignore=None):
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            if os.path.exists(d):
                shutil.rmtree(d)
            shutil.copytree(s, d, symlinks, ignore)
        else:
            shutil.copy2(s, d)


# obtain Pennsieve S3 URL for an existing metadata file
def returnFileURL(bf_object, item_id):

    file_details = bf_object._api._get(f"/packages/{str(item_id)}/view")
    file_id = file_details[0]["content"]["id"]
    file_url_info = bf_object._api._get(
        f"/packages/{str(item_id)}/files/{str(file_id)}"
    )


    return file_url_info["url"]


def recursive_item_path_create(folder, path):
    if "files" in folder.keys():
        for item in list(folder["files"]):
            if "folderpath" not in folder["files"][item]:
                folder["files"][item]["folderpath"] = path[:]

    if "folders" in folder.keys():
        for item in list(folder["folders"]):
            if "folderpath" not in folder["folders"][item]:
                folder["folders"][item]["folderpath"] = path[:]
                folder["folders"][item]["folderpath"].append(item)
            recursive_item_path_create(
                folder["folders"][item], folder["folders"][item]["folderpath"][:]
            )

    return


## import an existing local or Pennsieve dataset_description.xlsx file
def load_existing_DD_file(import_type, filepath):

    ### the following block of code converts .xlsx file into .csv for better performance from Pandas.
    ### Currently pandas' read_excel is super slow - could take minutes.
    # open given workbook
    # and store in excel object

    if import_type == "bf":
        try:

            DD_df = pd.read_excel(
                filepath, engine="openpyxl", usecols=column_check, header=0
            )

        except:
            raise Exception(
                "SODA cannot read this submission.xlsx file. If you are trying to retrieve a submission.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
            )

    else:
        excel = load_workbook(filepath)
        sheet = excel.active
        # writer object is created
        with tempfile.NamedTemporaryFile(mode="w", newline="", delete=False) as tf:
            col = csv.writer(tf)
            # writing the data in csv file
            for r in sheet.rows:
                col.writerow([cell.value for cell in r])

        DD_df = pd.DataFrame(
            pd.read_csv(tf.name, encoding="ISO-8859-1", usecols=column_check, header=0)
        )

    DD_df = DD_df.dropna(axis=0, how="all")
    DD_df = DD_df.replace(np.nan, "", regex=True)
    DD_df = DD_df.applymap(str)
    DD_df = DD_df.applymap(str.strip)

    basicInfoHeaders = [
        "Type",
        "Title",
        "Subtitle",
        "Keywords",
        "Number of subjects",
        "Number of samples",
    ]
    studyInfoHeaders = [
        "Study purpose",
        "Study data collection",
        "Study primary conclusion",
        "Study organ system",
        "Study approach",
        "Study technique",
        "Study collection title",
    ]
    contributorInfoHeaders = [
        "Contributor name",
        "Contributor ORCiD",
        "Contributor affiliation",
        "Contributor role",
    ]
    awardInfoHeaders = ["Funding", "Acknowledgments"]
    relatedInfoHeaders = [
        "Identifier description",
        "Relation type",
        "Identifier",
        "Identifier type",
    ]

    header_list = list(
        itertools.chain(
            basicInfoHeaders,
            studyInfoHeaders,
            contributorInfoHeaders,
            awardInfoHeaders,
            relatedInfoHeaders,
        )
    )

    # check if Metadata Element a.k.a Header column exists
    for key in ["Metadata element", "Description", "Example", "Value"]:
        if key not in DD_df:
            abort(400, 
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
            )

    if "Metadata element" not in DD_df:
        abort(400, 
            "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
        )

    else:
        for header_name in header_list:
            if header_name not in set(DD_df["Metadata element"]):
                abort(400, 
                    "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
                )

    if non_empty_1st_value := checkEmptyColumn(DD_df["Value"]):
        abort(400, 
            "At least 1 value is required to import an existing dataset_description file"
        )

    # drop Description and Examples columns
    DD_df = DD_df.drop(columns=["Description", "Example"])



    ## convert DD_df to array of arrays (a matrix)
    DD_matrix = DD_df.to_numpy().tolist()

    # divide DD_df into UI sections (Basic Info + Participant Info, Study Info, Contributor Info, Related Info)
    basicInfoSection = []
    studyInfoSection = []
    conInfoSection = []
    awardInfoSection = []
    relatedInfoSection = []

    for array in DD_matrix:
        if array[0] in basicInfoHeaders:
            basicInfoSection.append(array)
        if array[0] in studyInfoHeaders:
            studyInfoSection.append(array)
        if array[0] in contributorInfoHeaders:
            conInfoSection.append(array)
        if array[0] in awardInfoHeaders:
            awardInfoSection.append(array)
        if array[0] in relatedInfoHeaders:
            relatedInfoSection.append(array)

    return {
        "Basic information": basicInfoSection,
        "Study information": studyInfoSection,
        "Contributor information": transposeMatrix(conInfoSection),
        "Award information": awardInfoSection,
        "Related information": transposeMatrix(relatedInfoSection),
    }


def delete_manifest_dummy_folders(userpath_list):
    for userpath in userpath_list:
        shutil.rmtree(userpath) if isdir(userpath) else 0


def edit_bf_manifest_file(edit_action, manifest_type):
    if manifest_type == "bf":
        manifest_file_location = os.path.join(userpath, "SODA", "manifest_files")
    else:
        manifest_file_location = os.path.join(userpath, "SODA", "Manifest Files")

    if edit_action == "drop_empty_manifest_columns":
        drop_manifest_empty_columns(manifest_file_location)

    return 


    


def drop_manifest_empty_columns(manifest_file_location):
    global namespace_logger
    # read the manifest files in the manifest files folder
    high_level_folders = os.listdir(manifest_file_location)

    # go through each high level folder
    for high_level_folder in high_level_folders:
        # read the folder's excel file 
        manifest_df = pd.read_excel(
                        os.path.join(manifest_file_location, high_level_folder, "manifest.xlsx"), engine="openpyxl", usecols=column_check, header=0
                    )
        custom_columns = []

        # get the custom columns from the data frame
        SET_COLUMNS = ['filename', 'timestamp', 'description', 'file type', 'Additional Metadata']
        for column in manifest_df.columns: 
            if column not in SET_COLUMNS:
                custom_columns.append(column)


        # for each custom column delete the column if all values are null/empty
        manifest_dict = {x:manifest_df[x].values.tolist() for x in manifest_df}

        for column in custom_columns:
            if all([pd.isna(x) for x in manifest_dict[column]]):
                # remove the column from dict
                del manifest_dict[column]

        # convert the dict to a data frame
        edited_manifest_df = pd.DataFrame.from_dict(manifest_dict)

        # save the data frame to the manifest folder as an excel file
        save_location = os.path.join(manifest_file_location, high_level_folder, "manifest.xlsx")
        edited_manifest_df.to_excel(save_location, index=False)
        wb = load_workbook(save_location)
        ws = wb.active
        blueFill = PatternFill(
            start_color="9DC3E6", fill_type="solid"
        )
        greenFill = PatternFill(
            start_color="A8D08D", fill_type="solid"
        )
        yellowFill = PatternFill(
            start_color="FFD965", fill_type="solid"
        )
        ws['A1'].fill = blueFill
        ws['B1'].fill = greenFill
        ws['C1'].fill = greenFill
        ws['D1'].fill = greenFill
        ws['E1'].fill = yellowFill
        wb.save(save_location)

