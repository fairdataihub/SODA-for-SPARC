# -*- coding: utf-8 -*-

### Import required python modules

import platform
import os
from os.path import (
    isdir,
    join,
    exists,
    expanduser,
    getsize,
)
import pandas as pd
import csv
import shutil
import numpy as np
import json
from functools import partial
from utils import ( connect_pennsieve_client, authenticate_user_with_client, get_dataset_id, create_request_headers, column_check, returnFileURL, load_metadata_to_dataframe)
from permissions import has_edit_permissions
from collections import defaultdict
import requests
from errorHandlers import is_file_not_found_exception, is_invalid_file_exception, InvalidDeliverablesDocument
import time
from authentication import get_access_token


from string import ascii_uppercase
import itertools
import tempfile

from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font
from docx import Document

from flask import abort 
from constants import PENNSIEVE_URL





from manifest import update_existing_pennsieve_manifest_files, create_high_lvl_manifest_files_existing_ps_starting_point, recursive_item_path_create

from namespaces import NamespaceEnum, get_namespace_logger
namespace_logger = get_namespace_logger(NamespaceEnum.CURATE_DATASETS)

userpath = expanduser("~")
METADATA_UPLOAD_BF_PATH = join(userpath, "SODA", "METADATA")
TEMPLATE_PATH = ""
manifest_folder_path = join(userpath, "SODA", "manifest_files")

### Sets the TEMPLATE_PATH using SODA-for-SPARC's basepath so that the prepare_metadata section can find
### the templates stored in file_templates directory
def set_template_path(soda_base_path, soda_resources_path):
    """
    Sets the TEMPLATE_PATH using SODA-for-SPARC's basepath so that the prepare_metadata section can find
    the templates stored in file_templates directory.
    """
    global TEMPLATE_PATH


    # once pysoda has been packaged with pyinstaller
    # it creates an archive that slef extracts to an OS-specific temp directory.
    # Due to this we can no longer use a relative path from the pysoda directory to the file_templates folder.
    # When running in dev mode this also works
    TEMPLATE_PATH = join(soda_base_path, "file_templates")

    # only true if running SODA for SPARC as a build
    if platform.system() in ["Darwin", "Linux", "Windows"] and not exists(TEMPLATE_PATH):
        # we are in production and we need to use the Resources folder for the file_templates folder
        TEMPLATE_PATH = join(soda_resources_path, "file_templates")

    if not exists(TEMPLATE_PATH):
        TEMPLATE_PATH = join(soda_base_path, "file_templates")





### Import Data Deliverables document


def import_sparc_award(filepath):
    """
        Import SPARC Award from Data Deliverables document
    """
    doc = Document(filepath)
    ## Go through the paragraphs in the doc, and return the text of the first paragraph that contains the word "SPARC award #:"
    ## If no paragraph contains the word "SPARC award #:", then return an empty string
    sparc_award = ""
    for paragraph in doc.paragraphs:
        if "SPARC award #:" in paragraph.text:
            sparc_award = paragraph.text.split(":")[1].strip()
    return sparc_award


def import_milestone(filepath):
    """
        Import Data Deliverables document
    """
    doc = Document(filepath)
    try:
        table = doc.tables[0]
    except IndexError:
        raise InvalidDeliverablesDocument(
            "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: 'Related milestone, aim, or task', 'Description of data', and 'Expected date of completion'."
        )
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # headers will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def extract_milestone_info(datalist):
    milestone = defaultdict(list)
    milestone_key1 = "Related milestone, aim, or task"
    milestone_key2 = "Related milestone, aim or task"
    other_keys = ["Description of data", "Expected date of completion"]
    for row in datalist:
        if milestone_key1 in row:
            milestone_key = milestone_key1
        elif milestone_key2 in row:
            milestone_key = milestone_key2
        else:
            raise InvalidDeliverablesDocument(
                "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: Related milestone, aim, or task, Description of data, and Expected date of completion."
            )

        key = row[milestone_key]
        if key != "":
            milestone[key].append({key: row[key] for key in other_keys})

    return milestone







def subscriber_metadata(ps, events_dict):
    if events_dict["type"] == 1:
        fileid = events_dict["upload_status"].file_id
        total_bytes_to_upload = events_dict["upload_status"].total
        current_bytes_uploaded = events_dict["upload_status"].current
        if current_bytes_uploaded == total_bytes_to_upload and fileid != "":
            namespace_logger.info("File upload complete")
            ps.unsubscribe(10)

def upload_metadata_file(file_type, bfaccount, bfdataset, file_path, delete_after_upload):
    # check that the Pennsieve dataset is valid
    selected_dataset_id = get_dataset_id(bfdataset)

    # check that the user has permissions for uploading and modifying the dataset
    if not has_edit_permissions(get_access_token(), selected_dataset_id):
        abort(403, "You do not have permissions to edit this dataset.")
    headers = create_request_headers(get_access_token())
    # handle duplicates on Pennsieve: first, obtain the existing file ID
    r = requests.get(f"{PENNSIEVE_URL}/datasets/{selected_dataset_id}", headers=headers)
    r.raise_for_status()
    ds_items = r.json()
    # go through the content in the dataset and find the file ID of the file to be uploaded
    for item in ds_items["children"]:
        if item["content"]["name"] == file_type:
            item_id = item["content"]["id"]
            jsonfile = {
                "things": [item_id]
            }
            # then, delete it using Pennsieve method delete(id)\vf = Pennsieve()
            r = requests.post(f"{PENNSIEVE_URL}/data/delete",json=jsonfile, headers=headers)
            r.raise_for_status()
    try:
        ps = connect_pennsieve_client(bfaccount)
        authenticate_user_with_client(ps, bfaccount)
        # create a new manifest for the metadata file
        ps.use_dataset(selected_dataset_id)
        manifest = ps.manifest.create(file_path)
        m_id = manifest.manifest_id
    except Exception as e:
        error_message = "Could not create manifest file for this dataset"
        abort(500, error_message)
    
    # upload the manifest file
    try: 
        ps.manifest.upload(m_id)
        # create a subscriber function with ps attached so it can be used to unusbscribe
        subscriber_metadata_ps_client = partial(subscriber_metadata, ps)
        # subscribe for the upload to finish
        ps.subscribe(10, False, subscriber_metadata_ps_client)
    except Exception as e:
        namespace_logger.error("Error uploading dataset files")
        namespace_logger.error(e)
        raise Exception("The Pennsieve Agent has encountered an issue while uploading. Please retry the upload. If this issue persists please follow this <a target='_blank' rel='noopener noreferrer' href='https://docs.sodaforsparc.io/docs/how-to/how-to-reinstall-the-pennsieve-agent'> guide</a> on performing a full reinstallation of the Pennsieve Agent to fix the problem.")


    # before we can remove files we need to wait for all of the Agent's threads/subprocesses to finish
    # elsewise we get an error that the file is in use and therefore cannot be deleted
    time.sleep(5)

    # delete the local file that was created for the purpose of uploading to Pennsieve
    if delete_after_upload:
        os.remove(file_path)









### Prepare dataset-description file















# import an existing subjects/samples files from an excel file
def convert_subjects_samples_file_to_df(type, filepath, ui_fields, item_id=None, token=None):
    if item_id is not None: 
        subjects_df = load_metadata_to_dataframe(item_id, "excel", token, column_check, 0)
    else:
        subjects_df = pd.read_excel(
            filepath, engine="openpyxl", usecols=column_check, header=0
        )
    subjects_df = subjects_df.dropna(axis=0, how="all")
    subjects_df = subjects_df.replace(np.nan, "", regex=True)
    subjects_df = subjects_df.applymap(str)
    subjects_df.columns = map(str.lower, subjects_df.columns)

    if type == "subjects":
        if "subject id" not in list(subjects_df.columns.values):
            abort(
                400, 
                "The header 'subject id' is required to import an existing subjects file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/subjects.xlsx'>template</a> of the subjects file."
            )

        elif checkEmptyColumn(subjects_df["subject id"]):
            abort(
                400,
                "At least 1 'subject id' is required to import an existing subjects file"
            )

        templateHeaderList = subjectsTemplateHeaderList

    else:
        if "subject id" not in list(subjects_df.columns.values) or "sample id" not in list(subjects_df.columns.values):
            abort(
                400,
                "The headers 'subject id' and 'sample id' are required to import an existing samples file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/samples.xlsx'>template</a> of the samples file."
            )

        if checkEmptyColumn(subjects_df["sample id"]) or checkEmptyColumn(
            subjects_df["sample id"]
        ):
            abort(
                400,
                "At least 1 'subject id' and 'sample id' pair is required to import an existing samples file"
            )

        templateHeaderList = samplesTemplateHeaderList

    importedHeaderList = list(subjects_df.columns.values)

    transpose = []
    for header in templateHeaderList:
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    for header in importedHeaderList:

        if header.lower() in templateHeaderList:
            continue
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    sortMatrix = sortedSubjectsTableData(transpose, ui_fields)

    return {"sample_file_rows": transposeMatrix(sortMatrix)} if type in ["samples.xlsx", "samples"] else {"subject_file_rows": transposeMatrix(sortMatrix)}


## check if any whole column from Excel sheet is empty
def checkEmptyColumn(column):
    for element in column:
        if element:
            break
        return True
    return False



## load/import an existing local or Pennsieve submission.xlsx file
def load_existing_submission_file(filepath, item_id=None, token=None):

    try:
        if item_id is None:
            submission_data_frame = pd.read_excel(
                filepath, engine="openpyxl", usecols=column_check, header=0
            )
        else:
            submission_data_frame = load_metadata_to_dataframe(item_id, "excel", token, column_check, 0)

    except Exception as e:
        if is_file_not_found_exception(e):
            abort(400, "Local submission file not found")

        if is_invalid_file_exception(e):
            abort(400, "Local submission file is not in the correct format")

        raise Exception(
            "SODA cannot read this submission.xlsx file. If you are trying to retrieve a submission.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
        ) from e

    # drop rows with missing values, convert values to strings, and remove white spaces
    submission_data_frame = submission_data_frame.dropna(axis=0, how="all")
    submission_data_frame = submission_data_frame.replace(np.nan, "", regex=True)
    submission_data_frame = submission_data_frame.applymap(str)
    submission_data_frame = submission_data_frame.applymap(str.strip)

    basicColumns = ["Submission Item", "Definition", "Value"]
    basicHeaders = [
        "Consortium data standard",
        "Funding consortium",
        "Award number",
        "Milestone achieved",
        "Milestone completion date",
    ]

    #log the submission_data_frame
    ## normalize the entries to lowercase just for Version Exception check
    basicColumns = [x.lower() for x in basicColumns]
    submission_data_frame_lowercased = [x.lower() for x in submission_data_frame]
    for key in basicColumns:
        if key not in submission_data_frame_lowercased:
            abort(
                400,
                "The imported file columns are not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.1.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )

    basicHeaders = [x.lower() for x in basicHeaders]
    for header_name in basicHeaders:
        submissionItems = [x.lower() for x in submission_data_frame["Submission Item"]]
        if header_name not in set(submissionItems):
            abort(
                400,
                "The imported file headers are not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.1.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )


    consortium_data_standard = submission_data_frame["Value"][0]
    funding_consortium = submission_data_frame["Value"][1]
    award_number = submission_data_frame["Value"][2]
    milestones = [submission_data_frame["Value"][3]]

    for i in range(3, len(submission_data_frame.columns)):
        value = submission_data_frame[f"Value {str(i - 1)}"]
        milestones.append(value[3])

    milestone_completion_data = submission_data_frame["Value"][4] or ""

    return {
        "Consortium data standard": consortium_data_standard,
        "Funding consortium": funding_consortium,
        "Award number": award_number,
        "Milestone achieved": milestones,
        "Milestone completion date": milestone_completion_data,
    }


# import existing metadata files except Readme and Changes from Pennsieve
def import_ps_metadata_file(file_type, ui_fields, bfdataset):
    selected_dataset_id = get_dataset_id(bfdataset)

    r = requests.get(f"{PENNSIEVE_URL}/datasets/{selected_dataset_id}", headers=create_request_headers(get_access_token()))
    r.raise_for_status()

    ds_items = r.json()["children"]

    for i in ds_items:
        if i["content"]["name"] == file_type:
            item_id = i["content"]["id"]
            url = returnFileURL(get_access_token(), item_id)

            if file_type == "submission.xlsx":
                return load_existing_submission_file(url, item_id, get_access_token())

            elif file_type == "dataset_description.xlsx":
                # bf is the old signifier for pennsieve
                return load_existing_DD_file("bf", url, item_id, get_access_token())

            elif file_type == "subjects.xlsx":
                return convert_subjects_samples_file_to_df("subjects", url, ui_fields, item_id, get_access_token())

            elif file_type == "samples.xlsx":
                return convert_subjects_samples_file_to_df("samples", url, ui_fields, item_id, get_access_token())
            
            elif file_type == "code_description.xlsx":
                # Simply return true since we don't currently have a UI for code_description
                return True

    abort(400, 
        f"No {file_type} file was found at the root of the dataset provided."
    )


# import readme or changes file from Pennsieve
def import_ps_RC(bfdataset, file_type):
    file_type = file_type + ".txt"

    dataset_id = get_dataset_id(bfdataset)

    r = requests.get(f"{PENNSIEVE_URL}/datasets/{dataset_id}", headers=create_request_headers(get_access_token()))
    r.raise_for_status()

    items = r.json()

    for item in items["children"]:
        if item["content"]["name"] == file_type:
            item_id = item["content"]["id"]
            url = returnFileURL(get_access_token(), item_id)
            r = requests.get(url)
            return {"text": r.text}

    abort (400, f"No {file_type} file was found at the root of the dataset provided.")


# path to local SODA folder for saving manifest files
manifest_folder_path = join(userpath, "SODA", "manifest_files")
manifest_progress = {
    "total_manifest_files": 0,
    "manifest_files_uploaded": 0,
    "finished": False
}


def import_ps_manifest_file(soda_json_structure):
    # reset the progress tracking information
    global manifest_progress
    global manifest_folder_path
    manifest_progress["finished"] = False
    manifest_progress["total_manifest_files"] = 0
    manifest_progress["manifest_files_uploaded"] = 0

    high_level_folders = ["code", "derivative", "docs", "primary", "protocol", "source"]
    # convert the string into a json object/dictionary
    if(str(type(soda_json_structure)) == "<class 'str'>"):
        soda_json_structure = json.loads(soda_json_structure)
        
    dataset_structure = soda_json_structure["dataset-structure"]

    # get the count of the total number of high level folders in soda_json_structure
    for folder in list(dataset_structure["folders"]):
        if folder in high_level_folders and (dataset_structure["folders"][folder]["files"] == {} and dataset_structure["folders"][folder]["folders"] == {}):
            manifest_progress["total_manifest_files"] += 1

    # create the path to the dataset files and folders on Pennsieve and add them to the dataset structure stored in soda_json_structure
    recursive_item_path_create(dataset_structure, [])
    high_level_folders = ["code", "derivative", "docs", "primary", "protocol", "source"]

    # handle updating any existing manifest files on Pennsieve
    update_existing_pennsieve_manifest_files(soda_json_structure, high_level_folders, manifest_progress, manifest_folder_path)
    # create manifest files from scratch for any high level folders that don't have a manifest file on Pennsieve
    create_high_lvl_manifest_files_existing_ps_starting_point(soda_json_structure, manifest_folder_path, high_level_folders, manifest_progress)
    # finished with the manifest generation process
    manifest_progress["finished"] = True
    
    return {"message": "Finished"}


def manifest_creation_progress():
    """
        Monitors the progress of the manifest creation process.
    """
    global manifest_progress

    return {
        "manifest_files_uploaded": manifest_progress["manifest_files_uploaded"], 
        "total_manifest_files": manifest_progress["total_manifest_files"],
        "finished": manifest_progress["finished"]
    }


def copytree(src, dst, symlinks=False, ignore=None):
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            if os.path.exists(d):
                shutil.rmtree(d)
            shutil.copytree(s, d, symlinks, ignore)
        else:
            shutil.copy2(s, d)


## import an existing local or Pennsieve dataset_description.xlsx file
def load_existing_DD_file(import_type, filepath, item_id=None, token=None):

    ### the following block of code converts .xlsx file into .csv for better performance from Pandas.
    ### Currently pandas' read_excel is super slow - could take minutes.
    # open given workbook
    # and store in excel object

     # bf is the old signifier for pennsieve
    if import_type == "bf":
        try:
            DD_df = load_metadata_to_dataframe(item_id, "excel", token, column_check, 0)
        except Exception as e:
            namespace_logger.info(e)
            raise Exception from e (
                "SODA cannot read this dataset_description.xlsx file. If you are trying to retrieve a dataset_description.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
            )

    else:
        excel = load_workbook(filepath)
        sheet = excel.active
        # writer object is created
        with tempfile.NamedTemporaryFile(mode="w", newline="", delete=False) as tf:
            col = csv.writer(tf)
            # writing the data in csv file
            for r in sheet.rows:
                col.writerow([cell.value for cell in r])

        DD_df = pd.DataFrame(
            pd.read_csv(tf.name, encoding="ISO-8859-1", usecols=column_check, header=0)
        )

    # drop emtpy rows, convert values to strings, and remove white spaces
    DD_df = DD_df.dropna(axis=0, how="all")
    DD_df = DD_df.replace(np.nan, "", regex=True)
    DD_df = DD_df.applymap(str)
    DD_df = DD_df.applymap(str.strip)

    basicInfoHeaders = [
        "Type",
        "Title",
        "Subtitle",
        "Keywords",
        "Number of subjects",
        "Number of samples",
    ]
    studyInfoHeaders = [
        "Study purpose",
        "Study data collection",
        "Study primary conclusion",
        "Study organ system",
        "Study approach",
        "Study technique",
        "Study collection title",
    ]
    contributorInfoHeaders = [
        "Contributor name",
        "Contributor ORCiD",
        "Contributor affiliation",
        "Contributor role",
    ]
    awardInfoHeaders = ["Funding", "Acknowledgments"]
    relatedInfoHeaders = [
        "Identifier description",
        "Relation type",
        "Identifier",
        "Identifier type",
    ]

    header_list = list(
        itertools.chain(
            basicInfoHeaders,
            studyInfoHeaders,
            contributorInfoHeaders,
            awardInfoHeaders,
            relatedInfoHeaders,
        )
    )

    # check if Metadata Element a.k.a Header column exists
    for key in ["Metadata element", "Description", "Example", "Value"]:
        if key not in DD_df:
            abort(400, 
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
            )

    if "Metadata element" not in DD_df:
        abort(400, 
            "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
        )

    else:
        for header_name in header_list:
            if header_name not in set(DD_df["Metadata element"]):
                abort(400, 
                    "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' rel='noopener noreferrer' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
                )

    if non_empty_1st_value := checkEmptyColumn(DD_df["Value"]):
        abort(400, 
            "At least 1 value is required to import an existing dataset_description file"
        )

    # drop Description and Examples columns
    DD_df = DD_df.drop(columns=["Description", "Example"])

    ## convert DD_df to array of arrays (a matrix)
    DD_matrix = DD_df.to_numpy().tolist()

    # divide DD_df into UI sections (Basic Info + Participant Info, Study Info, Contributor Info, Related Info)
    basicInfoSection = []
    studyInfoSection = []
    conInfoSection = []
    awardInfoSection = []
    relatedInfoSection = []

    for array in DD_matrix:
        if array[0] in basicInfoHeaders:
            basicInfoSection.append(array)
        if array[0] in studyInfoHeaders:
            studyInfoSection.append(array)
        if array[0] in contributorInfoHeaders:
            conInfoSection.append(array)
        if array[0] in awardInfoHeaders:
            awardInfoSection.append(array)
        if array[0] in relatedInfoHeaders:
            relatedInfoSection.append(array)

    return {
        "Basic information": basicInfoSection,
        "Study information": studyInfoSection,
        "Contributor information": transposeMatrix(conInfoSection),
        "Award information": awardInfoSection,
        "Related information": transposeMatrix(relatedInfoSection),
    }

def delete_manifest_dummy_folders(userpath_list):
    """
        Delete local manifest folders when a user switches from PS to local in the standalone manifest generator without saving. 
    """
    for userpath in userpath_list:
        shutil.rmtree(userpath) if isdir(userpath) else 0


