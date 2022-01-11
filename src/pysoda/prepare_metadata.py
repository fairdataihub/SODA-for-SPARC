# -*- coding: utf-8 -*-

### Import required python modules

from gevent import monkey

monkey.patch_all()
import platform
import os
from os import listdir, stat, makedirs, mkdir, walk, remove, pardir
from os.path import (
    isdir,
    isfile,
    join,
    splitext,
    getmtime,
    basename,
    normpath,
    exists,
    expanduser,
    split,
    dirname,
    getsize,
    abspath,
)
import pandas as pd
import csv
import io
import time
from time import strftime, localtime
import shutil
from shutil import copy2
from configparser import ConfigParser
import numpy as np
from collections import defaultdict
import subprocess
from websocket import create_connection
import socket
import errno
import re
import gevent
from pennsieve import Pennsieve
from pennsieve.log import get_logger
from pennsieve.api.agent import agent_cmd
from pennsieve.api.agent import AgentError, check_port, socket_address
from urllib.request import urlopen
import json
import collections
from threading import Thread
import pathlib
import requests

from string import ascii_uppercase
import itertools
import tempfile

from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.styles.colors import Color
from docx import Document

from datetime import datetime, timezone

from Bio import Entrez

from pysoda import (
    bf_get_current_user_permission,
    agent_running,
)

userpath = expanduser("~")
METADATA_UPLOAD_BF_PATH = join(userpath, "SODA", "METADATA")
DEV_TEMPLATE_PATH = join(dirname(__file__), "..", "file_templates")
# once pysoda has been packaged with pyinstaller
# it becomes nested into the pysodadist/api directory
PROD_TEMPLATE_PATH = join(dirname(__file__), "..", "..", "file_templates")
TEMPLATE_PATH = DEV_TEMPLATE_PATH if exists(DEV_TEMPLATE_PATH) else PROD_TEMPLATE_PATH

# custom Exception class for when a DDD file is in an invalid form
class InvalidDeliverablesDocument(Exception):
    pass


### Import Data Deliverables document
def import_milestone(filepath):
    doc = Document(filepath)
    try:
        table = doc.tables[0]
    except IndexError:
        raise InvalidDeliverablesDocument(
            "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: 'Related milestone, aim, or task', 'Description of data', and 'Expected date of completion'."
        )
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # headers will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def extract_milestone_info(datalist):
    milestone = defaultdict(list)
    milestone_key1 = "Related milestone, aim, or task"
    milestone_key2 = "Related milestone, aim or task"
    other_keys = ["Description of data", "Expected date of completion"]
    for row in datalist:
        if milestone_key1 in row:
            milestone_key = milestone_key1
        elif milestone_key2 in row:
            milestone_key = milestone_key2
        else:
            raise InvalidDeliverablesDocument(
                "Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: Related milestone, aim, or task, Description of data, and Expected date of completion."
            )

        key = row[milestone_key]
        if key != "":
            milestone[key].append({key: row[key] for key in other_keys})

    return milestone


### Create submission file
def save_submission_file(upload_boolean, bfaccount, bfdataset, filepath, json_str):

    font_submission = Font(name="Calibri", size=14, bold=False)

    source = join(TEMPLATE_PATH, "submission.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "submission.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)

    # json array to python list
    val_arr = json.loads(json_str)
    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]
    # date_obj = datetime.strptime(val_arr[2], "%Y-%m")
    # date_new = date_obj.strftime("%m-%Y")
    for column, arr in zip(excel_columns(start_index=2), val_arr):
        ws1[column + "2"] = arr["award"]
        ws1[column + "3"] = arr["milestone"]
        ws1[column + "4"] = arr["date"]

        ws1[column + "2"].font = font_submission
        ws1[column + "3"].font = font_submission
        ws1[column + "4"].font = font_submission

    rename_headers(ws1, len(val_arr), 2)

    wb.save(destination)

    # calculate the size of the metadata file
    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file("submission.xlsx", bfaccount, bfdataset, destination)

    return size


# this function saves and uploads the README/CHANGES to Pennsieve, just when users choose to generate onto Pennsieve
## (not used for generating locally)
def upload_RC_file(text_string, file_type, bfaccount, bfdataset):

    file_path = join(METADATA_UPLOAD_BF_PATH, file_type)

    with open(file_path, "w") as f:
        f.write(text_string)

    size = getsize(file_path)

    upload_metadata_file(file_type, bfaccount, bfdataset, file_path)

    return size, file_path


def upload_metadata_file(file_type, bfaccount, bfdataset, file_path):
    ## check if agent is running in the background
    agent_running()

    bf = Pennsieve(bfaccount)

    # check that the Pennsieve dataset is valid
    try:
        myds = bf.get_dataset(bfdataset)

    except Exception as e:
        error = "Error: Please select a valid Pennsieve dataset"
        raise Exception(error)

    else:
        # check that the user has permissions for uploading and modifying the dataset
        role = bf_get_current_user_permission(bf, myds)
        if role not in ["owner", "manager", "editor"]:
            error = "Error: You don't have permissions for uploading to this Pennsieve dataset."
            raise Exception(error)

        # handle duplicates on Pennsieve: first, obtain the existing file ID
        for i in range(len(myds.items)):

            if myds.items[i].name == file_type:

                item_id = myds.items[i].id

                # then, delete it using Pennsieve method delete(id)
                bf.delete(item_id)

        myds.upload(file_path)
        # delete the local file that was created for the purpose of uploading to Pennsieve
        os.remove(file_path)


def excel_columns(start_index=0):
    """
    NOTE: does not support more than 699 contributors/links
    """
    single_letter = list(ascii_uppercase[start_index:])
    two_letter = [a + b for a, b in itertools.product(ascii_uppercase, ascii_uppercase)]
    return single_letter + two_letter


def rename_headers(workbook, max_len, start_index):
    """
    Rename header columns if values exceed 3. Change Additional Values to Value 4, 5,...
    """

    columns_list = excel_columns(start_index=start_index)
    if max_len >= start_index:

        workbook[columns_list[0] + "1"] = "Value"

        for i, column in zip(range(2, max_len + 1), columns_list[1:]):

            workbook[column + "1"] = "Value " + str(i)
            cell = workbook[column + "1"]

            blueFill = PatternFill(
                start_color="9CC2E5", end_color="9CC2E5", fill_type="solid"
            )

            font = Font(bold=True)
            cell.fill = blueFill
            cell.font = font

    else:

        delete_range = len(columns_list) - max_len
        workbook.delete_cols(4 + max_len, delete_range)


def grayout_subheaders(workbook, max_len, start_index):
    """
    Gray out sub-header rows for values exceeding 3 (SDS2.0).
    """
    headers_list = ["4", "10", "18", "23", "28"]
    columns_list = excel_columns(start_index=start_index)

    for i, column in zip(range(2, max_len + 1), columns_list[1:]):

        for no in headers_list:
            cell = workbook[column + no]
            fillColor("B2B2B2", cell)


def grayout_single_value_rows(workbook, max_len, start_index):
    """
    Gray out rows where only single values are allowed. Row number: 2, 3, 5, 6, 9, 11, 12, 13, 17, 29, 30
    """

    columns_list = excel_columns(start_index=start_index)
    row_list = ["2", "3", "5", "6", "9", "11", "12", "13", "17", "29", "30"]
    for i, column in zip(range(2, max_len + 1), columns_list[1:]):

        for no in row_list:
            cell = workbook[column + no]
            fillColor("CCCCCC", cell)


def fillColor(color, cell):
    colorFill = PatternFill(start_color=color, end_color=color, fill_type="solid")

    cell.fill = colorFill


### Prepare dataset-description file


def populate_dataset_info(workbook, val_obj):
    ## name, description, type, samples, subjects
    workbook["D5"] = val_obj["name"]
    workbook["D6"] = val_obj["description"]
    workbook["D3"] = val_obj["type"]
    workbook["D29"] = val_obj["number of subjects"]
    workbook["D30"] = val_obj["number of samples"]

    ## keywords
    for i, column in zip(range(len(val_obj["keywords"])), excel_columns(start_index=3)):
        workbook[column + "7"] = val_obj["keywords"][i]

    return val_obj["keywords"]


def populate_study_info(workbook, val_obj):
    workbook["D11"] = val_obj["study purpose"]
    workbook["D12"] = val_obj["study data collection"]
    workbook["D13"] = val_obj["study primary conclusion"]
    workbook["D17"] = val_obj["study collection title"]

    ## study organ system
    for i, column in zip(
        range(len(val_obj["study organ system"])), excel_columns(start_index=3)
    ):
        workbook[column + "14"] = val_obj["study organ system"][i]
    ## study approach
    for i, column in zip(
        range(len(val_obj["study approach"])), excel_columns(start_index=3)
    ):
        workbook[column + "15"] = val_obj["study approach"][i]
    ## study technique
    for i, column in zip(
        range(len(val_obj["study technique"])), excel_columns(start_index=3)
    ):
        workbook[column + "16"] = val_obj["study technique"][i]

    return max(
        len(val_obj["study organ system"]),
        len(val_obj["study approach"]),
        len(val_obj["study technique"]),
    )


def populate_contributor_info(workbook, val_array):
    ## award info
    for i, column in zip(
        range(len(val_array["funding"])), excel_columns(start_index=3)
    ):
        workbook[column + "8"] = val_array["funding"][i]

    ### Acknowledgments
    workbook["D9"] = val_array["acknowledgment"]

    ### Contributors
    for contributor, column in zip(
        val_array["contributors"], excel_columns(start_index=3)
    ):
        workbook[column + "19"] = contributor["conName"]
        workbook[column + "20"] = contributor["conID"]
        workbook[column + "21"] = contributor["conAffliation"]
        workbook[column + "22"] = contributor["conRole"]

    return [val_array["funding"], val_array["contributors"]]


def populate_related_info(workbook, val_array):
    ## related links including protocols

    for i, column in zip(range(len(val_array)), excel_columns(start_index=3)):
        workbook[column + "24"] = val_array[i]["description"]
        workbook[column + "25"] = val_array[i]["relation"]
        workbook[column + "26"] = val_array[i]["link"]
        workbook[column + "27"] = val_array[i]["type"]

    return len(val_array)


### generate the dataset_description file


def save_ds_description_file(
    upload_boolean,
    bfaccount,
    bfdataset,
    filepath,
    dataset_str,
    study_str,
    con_str,
    related_info_str,
):
    source = join(TEMPLATE_PATH, "dataset_description.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "dataset_description.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)

    # json array to python list
    val_obj_study = json.loads(study_str)
    val_obj_ds = json.loads(dataset_str)
    val_arr_con = json.loads(con_str)
    val_arr_related_info = json.loads(related_info_str)

    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    ws1["D22"] = ""
    ws1["E22"] = ""
    ws1["D24"] = ""
    ws1["E24"] = ""
    ws1["D25"] = ""
    ws1["E25"] = ""

    keyword_array = populate_dataset_info(ws1, val_obj_ds)

    study_array_len = populate_study_info(ws1, val_obj_study)

    (funding_array, contributor_role_array) = populate_contributor_info(
        ws1, val_arr_con
    )

    related_info_len = populate_related_info(ws1, val_arr_related_info)

    # keywords length
    keyword_len = len(keyword_array)

    # contributors length
    no_contributors = len(contributor_role_array)

    # funding = SPARC award + other funding sources
    funding_len = len(funding_array)

    # obtain length for formatting compliance purpose
    max_len = max(
        keyword_len, funding_len, no_contributors, related_info_len, study_array_len
    )

    rename_headers(ws1, max_len, 3)
    grayout_subheaders(ws1, max_len, 3)
    grayout_single_value_rows(ws1, max_len, 3)

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file(
            "dataset_description.xlsx", bfaccount, bfdataset, destination
        )

    return size


subjectsTemplateHeaderList = [
    "subject id",
    "pool id",
    "subject experimental group",
    "age",
    "sex",
    "species",
    "strain",
    "rrid for strain",
    "age category",
    "also in dataset",
    "member of",
    "laboratory internal id",
    "date of birth",
    "age range (min)",
    "age range (max)",
    "body mass",
    "genotype",
    "phenotype",
    "handedness",
    "reference atlas",
    "experimental log file path",
    "experiment date",
    "disease or disorder",
    "intervention",
    "disease model",
    "protocol title",
    "protocol url or doi",
]
samplesTemplateHeaderList = [
    "sample id",
    "subject id",
    "was derived from",
    "pool id",
    "sample experimental group",
    "sample type",
    "sample anatomical location",
    "also in dataset",
    "member of",
    "laboratory internal id",
    "date of derivation",
    "experimental log file path",
    "reference atlas",
    "pathology",
    "laterality",
    "cell type",
    "plane of section",
    "protocol title",
    "protocol url or doi",
]


def save_subjects_file(upload_boolean, bfaccount, bfdataset, filepath, datastructure):

    source = join(TEMPLATE_PATH, "subjects.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "subjects.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)
    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:11]
    optionalFields = transposeDatastructure[11:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = subjectsTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(
            np.concatenate((sortMatrix, refinedOptionalFields))
        )
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)
    #
    # # 1. delete rows using delete_rows(index, amount=2) -- description and example rows
    # ws1.delete_rows(2, 2)
    # delete all optional columns first (from the template)
    ws1.delete_cols(12, 18)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(
        start_color="FFD965", end_color="FFD965", fill_type="solid"
    )

    # gevent.sleep(0)
    for column, header in zip(
        excel_columns(start_index=11), refinedDatastructure[0][11:headers_no]
    ):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name="Calibri")

    # gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            if refinedDatastructure[i][j]:
                ws1[cell] = refinedDatastructure[i][j]
            else:
                ws1[cell] = ""
            ws1[cell].font = Font(bold=False, size=11, name="Arial")

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, then call upload function and then delete the destination path
    if upload_boolean:
        upload_metadata_file("subjects.xlsx", bfaccount, bfdataset, destination)

    return size


def save_samples_file(upload_boolean, bfaccount, bfdataset, filepath, datastructure):
    source = join(TEMPLATE_PATH, "samples.xlsx")

    if upload_boolean:
        destination = join(METADATA_UPLOAD_BF_PATH, "samples.xlsx")

    else:
        destination = filepath

    shutil.copyfile(source, destination)

    wb = load_workbook(destination)
    ws1 = wb["Sheet1"]

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:9]
    optionalFields = transposeDatastructure[9:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = samplesTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(
            np.concatenate((sortMatrix, refinedOptionalFields))
        )
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)

    ws1.delete_cols(10, 15)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(
        start_color="FFD965", end_color="FFD965", fill_type="solid"
    )
    # gevent.sleep(0)
    for column, header in zip(
        excel_columns(start_index=9), refinedDatastructure[0][9:headers_no]
    ):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name="Calibri")

    # gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            if refinedDatastructure[i][j]:
                ws1[cell] = refinedDatastructure[i][j]
            else:
                ws1[cell] = ""
            ws1[cell].font = Font(bold=False, size=11, name="Arial")

    wb.save(destination)

    size = getsize(destination)

    ## if generating directly on Pennsieve, call upload function
    if upload_boolean:
        upload_metadata_file("samples.xlsx", bfaccount, bfdataset, destination)

    return size


# check for non-empty fields (cells)
def column_check(x):
    if "unnamed" in x.lower():
        return False
    return True


# import an existing subjects/samples files from an excel file
def convert_subjects_samples_file_to_df(type, filepath, ui_fields):

    subjects_df = pd.read_excel(
        filepath, engine="openpyxl", usecols=column_check, header=0
    )
    subjects_df = subjects_df.dropna(axis=0, how="all")
    subjects_df = subjects_df.replace(np.nan, "", regex=True)
    subjects_df = subjects_df.applymap(str)
    subjects_df.columns = map(str.lower, subjects_df.columns)

    if type == "subjects":
        if "subject id" not in list(subjects_df.columns.values):
            raise Exception(
                "The header 'subject id' is required to import an existing subjects file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/subjects.xlsx'>template</a> of the subjects file."
            )

        else:
            if checkEmptyColumn(subjects_df["subject id"]):
                raise Exception(
                    "At least 1 'subject id' is required to import an existing subjects file"
                )

        templateHeaderList = subjectsTemplateHeaderList

    else:
        if "subject id" not in list(
            subjects_df.columns.values
        ) or "sample id" not in list(subjects_df.columns.values):
            raise Exception(
                "The headers 'subject id' and 'sample id' are required to import an existing samples file. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/samples.xlsx'>template</a> of the samples file."
            )

        else:
            if checkEmptyColumn(subjects_df["sample id"]) or checkEmptyColumn(
                subjects_df["sample id"]
            ):
                raise Exception(
                    "At least 1 'subject id' and 'sample id' pair is required to import an existing samples file"
                )

        templateHeaderList = samplesTemplateHeaderList

    importedHeaderList = list(subjects_df.columns.values)

    transpose = []
    for header in templateHeaderList:
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    for header in importedHeaderList:
        if header.lower() in templateHeaderList:
            continue
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""] * len(subjects_df))
        transpose.append(column)

    sortMatrix = sortedSubjectsTableData(transpose, ui_fields)

    return transposeMatrix(sortMatrix)


def checkEmptyColumn(column):
    for element in column:
        if element:
            break
        return True
    return False


# needed to sort subjects and samples table data to match the UI fields
def sortedSubjectsTableData(matrix, fields):
    sortedMatrix = []
    customHeaderMatrix = []

    for field in fields:
        for column in matrix:
            if column[0].lower() == field:
                sortedMatrix.append(column)
                break

    for column in matrix:
        if column[0].lower() not in fields:
            customHeaderMatrix.append(column)

    if len(customHeaderMatrix) > 0:
        npArray = np.concatenate((sortedMatrix, customHeaderMatrix)).tolist()
    else:
        npArray = sortedMatrix
    return npArray


# transpose a matrix (array of arrays)
def transposeMatrix(matrix):
    return [[matrix[j][i] for j in range(len(matrix))] for i in range(len(matrix[0]))]


# helper function to process custom fields (users add and name them) for subjects and samples files
def processMetadataCustomFields(matrix):
    refined_matrix = []
    for column in matrix:
        if any(column[1:]):
            refined_matrix.append(column)

    return refined_matrix


# use Entrez library to load the scientific name for a species
def load_taxonomy_species(animalList):
    animalDict = {}
    for animal in animalList:
        handle = Entrez.esearch(db="Taxonomy", term=animal)
        record = Entrez.read(handle)
        if len(record["IdList"]) > 0:
            id = record["IdList"][0]
            handle = Entrez.efetch(db="Taxonomy", id=id)
            result = Entrez.read(handle)
            animalDict[animal] = {
                "ScientificName": result[0]["ScientificName"],
                "OtherNames": result[0]["OtherNames"],
            }

    return animalDict


## check if any whole column from Excel sheet is empty
def checkEmptyColumn(column):
    for element in column:
        if element:
            break
        return True
    return False


## load/import an existing local or Pennsieve submission.xlsx file
def load_existing_submission_file(filepath):

    ## TODO: check and read csv
    try:
        DD_df = pd.read_excel(
            filepath, engine="openpyxl", usecols=column_check, header=0
        )

    except:
        raise Exception(
            "SODA cannot read this submission.xlsx file. If you are trying to retrieve a submission.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
        )

    DD_df = DD_df.dropna(axis=0, how="all")
    DD_df = DD_df.replace(np.nan, "", regex=True)
    DD_df = DD_df.applymap(str)
    DD_df = DD_df.applymap(str.strip)

    basicColumns = ["Submission Item", "Definition", "Value"]
    basicHeaders = [
        "SPARC Award number",
        "Milestone achieved",
        "Milestone completion date",
    ]
    ## normalize the entries to lowercase just for Version Exception check
    basicColumns = [x.lower() for x in basicColumns]
    basicHeaders = [x.lower() for x in basicHeaders]
    DD_df_lower = [x.lower() for x in DD_df]

    for key in basicColumns:
        if key not in DD_df_lower:
            raise Exception(
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )

    for header_name in basicHeaders:
        submissionItems = [x.lower() for x in DD_df["Submission Item"]]
        if header_name not in set(submissionItems):
            raise Exception(
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/submission.xlsx'>template</a> of the submission."
            )

    awardNo = DD_df["Value"][0]
    milestones = [DD_df["Value"][1]]

    for i in range(3, len(DD_df.columns)):
        value = DD_df["Value " + str(i - 1)]
        milestones.append(value[1])

    if DD_df["Value"][2]:
        date = DD_df["Value"][2]

    else:
        date = ""

    return {
        "SPARC Award number": awardNo,
        "Milestone achieved": milestones,
        "Milestone completion date": date,
    }


# import existing metadata files except Readme and Changes from Pennsieve
def import_bf_metadata_file(file_type, ui_fields, bfaccount, bfdataset):
    bf = Pennsieve(bfaccount)
    myds = bf.get_dataset(bfdataset)

    for i in range(len(myds.items)):

        if myds.items[i].name == file_type:

            item_id = myds.items[i].id
            url = returnFileURL(bf, item_id)

            if file_type == "submission.xlsx":
                return load_existing_submission_file(url)

            elif file_type == "dataset_description.xlsx":
                return load_existing_DD_file("bf", url)

            elif file_type == "subjects.xlsx":
                return convert_subjects_samples_file_to_df("subjects", url, ui_fields)

            elif file_type == "samples.xlsx":
                return convert_subjects_samples_file_to_df("samples", url, ui_fields)

    raise Exception(
        f"No {file_type} file was found at the root of the dataset provided."
    )


# import readme or changes file from Pennsieve
def import_bf_RC(bfaccount, bfdataset, file_type):
    bf = Pennsieve(bfaccount)
    myds = bf.get_dataset(bfdataset)

    for i in range(len(myds.items)):

        if myds.items[i].name == file_type:

            item_id = myds.items[i].id
            url = returnFileURL(bf, item_id)

            response = requests.get(url)
            data = response.text

            return data

    raise Exception(
        f"No {file_type} file was found at the root of the dataset provided."
    )


# obtain Pennsieve S3 URL for an existing metadata file
def returnFileURL(bf_object, item_id):

    file_details = bf_object._api._get("/packages/" + str(item_id) + "/view")
    file_id = file_details[0]["content"]["id"]
    file_url_info = bf_object._api._get(
        "/packages/" + str(item_id) + "/files/" + str(file_id)
    )

    return file_url_info["url"]


## import an existing local or Pennsieve dataset_description.xlsx file
def load_existing_DD_file(import_type, filepath):

    ### the following block of code converts .xlsx file into .csv for better performance from Pandas.
    ### Currently pandas' read_excel is super slow - could take minutes.
    # open given workbook
    # and store in excel object

    if import_type == "bf":
        try:

            DD_df = pd.read_excel(
                filepath, engine="openpyxl", usecols=column_check, header=0
            )

        except:
            raise Exception(
                "SODA cannot read this submission.xlsx file. If you are trying to retrieve a submission.xlsx file from Pennsieve, please make sure you are signed in with your Pennsieve account on SODA."
            )

    else:
        excel = load_workbook(filepath)
        sheet = excel.active
        # writer object is created
        with tempfile.NamedTemporaryFile(mode="w", newline="", delete=False) as tf:
            col = csv.writer(tf)
            # writing the data in csv file
            for r in sheet.rows:
                col.writerow([cell.value for cell in r])

        DD_df = pd.DataFrame(
            pd.read_csv(tf.name, encoding="ISO-8859-1", usecols=column_check, header=0)
        )

    DD_df = DD_df.dropna(axis=0, how="all")
    DD_df = DD_df.replace(np.nan, "", regex=True)
    DD_df = DD_df.applymap(str)
    DD_df = DD_df.applymap(str.strip)

    basicInfoHeaders = [
        "Type",
        "Title",
        "Subtitle",
        "Keywords",
        "Number of subjects",
        "Number of samples",
    ]
    studyInfoHeaders = [
        "Study purpose",
        "Study data collection",
        "Study primary conclusion",
        "Study organ system",
        "Study approach",
        "Study technique",
        "Study collection title",
    ]
    contributorInfoHeaders = [
        "Contributor name",
        "Contributor ORCiD",
        "Contributor affiliation",
        "Contributor role",
    ]
    awardInfoHeaders = ["Funding", "Acknowledgments"]
    relatedInfoHeaders = [
        "Identifier description",
        "Relation type",
        "Identifier",
        "Identifier type",
    ]

    header_list = list(
        itertools.chain(
            basicInfoHeaders,
            studyInfoHeaders,
            contributorInfoHeaders,
            awardInfoHeaders,
            relatedInfoHeaders,
        )
    )

    # check if Metadata Element a.k.a Header column exists
    for key in ["Metadata element", "Description", "Example", "Value"]:
        if key not in DD_df:
            raise Exception(
                "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
            )

    if not "Metadata element" in DD_df:
        raise Exception(
            "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
        )

    else:
        for header_name in header_list:
            if header_name not in set(DD_df["Metadata element"]):
                raise Exception(
                    "The imported file is not in the correct format. Please refer to the new SPARC Dataset Structure (SDS) 2.0.0 <a target='_blank' href='https://github.com/SciCrunch/sparc-curation/blob/master/resources/DatasetTemplate/dataset_description.xlsx'>template</a> of the dataset_description."
                )

    # check for at least 1 value is included
    non_empty_1st_value = checkEmptyColumn(DD_df["Value"])
    if non_empty_1st_value:
        raise Exception(
            "At least 1 value is required to import an existing dataset_description file"
        )

    # drop Description and Examples columns
    DD_df = DD_df.drop(columns=["Description", "Example"])

    ## convert DD_df to array of arrays (a matrix)
    DD_matrix = DD_df.to_numpy().tolist()

    # divide DD_df into UI sections (Basic Info + Participant Info, Study Info, Contributor Info, Related Info)
    basicInfoSection = []
    studyInfoSection = []
    conInfoSection = []
    awardInfoSection = []
    relatedInfoSection = []

    for array in DD_matrix:
        if array[0] in basicInfoHeaders:
            basicInfoSection.append(array)
        if array[0] in studyInfoHeaders:
            studyInfoSection.append(array)
        if array[0] in contributorInfoHeaders:
            conInfoSection.append(array)
        if array[0] in awardInfoHeaders:
            awardInfoSection.append(array)
        if array[0] in relatedInfoHeaders:
            relatedInfoSection.append(array)

    transformedObj = {
        "Basic information": basicInfoSection,
        "Study information": studyInfoSection,
        "Contributor information": transposeMatrix(conInfoSection),
        "Award information": awardInfoSection,
        "Related information": transposeMatrix(relatedInfoSection),
    }

    return transformedObj
